# -*- coding: utf-8 -*-

import re
import os
import zs_py as zp

from paper_structure import labels, type_names
from paper_structure import Paper
# from paper_parse import readClosing, readPartA, readPartB, readTsl, readWriting

# from docx_wrapper import DocxWrapper
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from nltk.corpus import wordnet as wn

circled_number = ('① ', '② ' , '③ ', '④ ' , '⑤ ' , '⑥ ' ,'⑦ ', '⑧ ' , '⑨ ' , '⑩ ' , '⑪ ' , '⑫ ' , '⑬ ' , '⑭ ' , '⑮ ')
circled_number_2 = ('❶ ', '❷ ' , '❸ ' , '❹ ' , '❺ ' , '❻ ' , '❼ ' , '❽ ' , '❾ ' , '➓ ' , '⓫ ' , '⓬ ' , '⓭ ' , '⓮ ' , '⓯ ' )

#---------------------------------------------------------------------------------------#
#                             Paper subs
#---------------------------------------------------------------------------------------#

import yaml

def load_from_yaml(file):
    with open(file, "r") as f:
        try:
            obj =  yaml.load(f, Loader=yaml.FullLoader)
            # obj =  yaml.safe_load(f)
        except yaml.YAMLError as exc:
            print(exc)
    return obj
    
class NoAliasDumper(yaml.Dumper):
    def ignore_aliases(self, data):
        return True
        
def paper_load_from_files(paper_files):
    paper = Paper()

    paper.closing = load_from_yaml( paper_files['Closing'] )
    paper.closing.split_sentence()
    # paper.closing.parse_question_position()
    
    paper.partA = [   load_from_yaml( file ) for file in [  paper_files['PartA1'], paper_files['PartA2'], paper_files['PartA3'], paper_files['PartA4']  ]   ]
    for i in range(0, 4):
        paper.partA[i].split_sentence()
        # paper.partA[i].parse_question_position()
    
    paper.partB = load_from_yaml( paper_files['PartB'] )
    paper.partB.split_sentence()
    paper.partB.shuffle(shuff_str=paper.partB.opinions_order)

    paper.tsl = load_from_yaml( paper_files['Trans'] )
    paper.tsl.split_sentence()

    paper.writing = load_from_yaml( paper_files['Writing'] )
    
    return paper

def paper_write_head_page(dw, paper_index):
    dw.write_paragraph('绝密★启用前', space_after=10, bold=True, font_size=10)
    dw.write_paragraph('', space_after=10)
    dw.write_paragraph('', space_after=10)
    dw.write_paragraph('2021 年全国硕士研究生入学统一考试', bold=True, align='CENTER', space_after=10, font_size=20)
    dw.write_paragraph('', space_after=10)
    dw.write_paragraph('管理类专业硕士学位联考', bold=True, align='CENTER', space_after=10, font_size=22)
    dw.write_paragraph('', space_after=10)
    dw.write_paragraph("英语（二）试卷" + paper_index, bold=True, align='CENTER', space_after=10, font_size=22)
    dw.write_paragraph('', space_after=10)

    dw.write_paragraph('考生须知', align='CENTER', space_after=10, font_size=14)
    dw.write_paragraph('1．考生必须严格遵守各项考场规则。', align='LEFT', space_after=10, font_size=10.5, left_indent=15, first_line_indent=-15)
    dw.write_paragraph('2．答题前，考生将答题卡上的“姓名”、“考生编号”等信息填写清楚，并与准考证上的一致。', align='LEFT', space_after=10, font_size=10.5, left_indent=15, first_line_indent=-15)
    dw.write_paragraph('3．选择题的答案须用2B铅笔填涂在答题卡上，其它笔填涂的或做在试卷或其它类型答题卡上的答案无效。', align='LEFT', space_after=10, font_size=10.5, left_indent=15, first_line_indent=-15)
    dw.write_paragraph('4．其他题一律用蓝色或黑色钢笔或圆珠笔在答题纸上按规定要求作答，凡做在试卷上或未做在指定位置的答案无效。', align='LEFT', space_after=10, font_size=10.5, left_indent=15, first_line_indent=-15)
    dw.write_paragraph('5．交卷时，请配合监考人员验收，并请监考人员在准考证相应位置签字（作为考生交卷的凭据）。否则，所产生的一切后果由考生自负。', align='LEFT', space_after=10, font_size=10.5, left_indent=15, first_line_indent=-15)

    dw.write_paragraph('', space_after=10)
    dw.write_paragraph('', space_after=10)

    #dw.write_paragraph('姓名：                      听课证号：', align='LEFT', font_size=10.5, left_indent=21, first_line_indent=50)

    dw.add_page_break()
    
def paper_write_closing(dw, closing):
    dw.write_paragraph('Section I  Use of English', font_size=14, align='CENTER')
    dw.write_paragraph('')
    
    dw.write_paragraph('Directions:', bold=True, space_after=5, first_line_indent=0)
    dw.write_paragraph('Read the following text. Choose the best word(s) for each numbered blank and mark A, B, C or D on the ANSWER SHEET. (10 points)', first_line_indent=0)
    dw.write_paragraph('')
    
    for para in closing.article_st:
        #dw.write_paragraph(trim(para))
        #[dw.write_paragraph(stce) for stce in para]
        text = ' '.join(para)
        dw.write_paragraph(text)
    
    dw.add_page_break()
    
    for qc_idx in range(1, 21):
        p = dw.write_paragraph(str(qc_idx) + '. ')
        p.add_run('[A] ' + closing.Q[qc_idx-1].A[0] + '\t')
        p.add_run('[B] ' + closing.Q[qc_idx-1].B[0] + '\t')
        p.add_run('[C] ' + closing.Q[qc_idx-1].C[0] + '\t')
        p.add_run('[D] ' + closing.Q[qc_idx-1].D[0])
    # dw.add_page_break()
        
def paper_write_partA(dw, partA):
    dw.write_paragraph('Section II  Reading Comprehension', font_size=14, align='CENTER')
    #dw.write_paragraph('')
    dw.write_paragraph('Part A', bold=True, first_line_indent=0, space_after=8)
    dw.write_paragraph('Directions:', bold=True, space_after=5, first_line_indent=0)
    dw.write_paragraph('Read the following four texts. Answer the questions after each text by choosing A, B, C or D. Mark your answers on the ANSWER SHEET. (40 points)', first_line_indent=0)
    #dw.write_paragraph('')
    dw.add_page_break()

    qa_index = 21
    for text_i in range(0, 4):
        try:
            p = dw.write_paragraph('Text '+ str(text_i+1), space_after=20, align='CENTER', bold=True)
            #p.paragraph_format.space_before = Pt(0)
            
            for para in partA[text_i].article_st:
                #dw.write_paragraph(trim(para))
                text = ' '.join(para)
                dw.write_paragraph(text)
            #dw.add_page_break()  # PAGE BREAK AFTER PART AN ARTICLE
            dw.add_page_break()
            
            for question_i in range(0, 5):
                dw.write_paragraph(str(qa_index) + ". " + partA[text_i].Q[question_i].question[0], left_indent=21, first_line_indent=-21)
                qa_index += 1
                dw.write_paragraph("[A] " + partA[text_i].Q[question_i].A[0], left_indent=0, first_line_indent=15)
                dw.write_paragraph("[B] " + partA[text_i].Q[question_i].B[0], left_indent=0, first_line_indent=15)
                dw.write_paragraph("[C] " + partA[text_i].Q[question_i].C[0], left_indent=0, first_line_indent=15)
                dw.write_paragraph("[D] " + partA[text_i].Q[question_i].D[0], left_indent=0, first_line_indent=15)
                dw.write_paragraph('')
                #dw.write_paragraph(document, "answer: ", partA[text_i].Q[question_i].answer)
        except Exception as e:
            print(e)
            print('paper_write_partA(): error at Text ', text_i+1, ',question ', question_i+1)
        dw.add_page_break()  # PAGE BREAK AFTER PART A QUESTIONS

def paper_write_partB(dw, partB):
    dw.write_paragraph('Part B', bold=True, first_line_indent=0, space_after=8)
    dw.write_paragraph('Directions:', bold=True, space_after=5, first_line_indent=0)
    if partB.meta.test == 'NEEP-II-B1':
        dw.write_paragraph('Read the following text and match each of the numbered items in the left column to its corresponding information in the right column. There are two extra choices in the right column. Mark your answers on the ANSWER SHEET. (10 points)', first_line_indent=0)

        for para in partB.article_st:
            #dw.write_paragraph(trim(para))
            text = ' '.join(para)
            dw.write_paragraph(text)
        #document.add_page_break()  # PAGE BREAK AFTER PART B ARTICLE

        dw.write_paragraph('')
        table = dw.add_table(rows=7, cols=2, style='Table Grid')
        qb_index = 41
        for char_i in range(1, 6):
            table.rows[char_i].cells[0].text = str(qb_index) + ". " +partB.characters[char_i-1]
            table.rows[char_i].cells[0].left_indent = 21
            table.rows[char_i].cells[0].first_line_indent = -21
            qb_index += 1

        for char_i in range(0, 7):
            table.rows[char_i].cells[1].text = "[" + labels[char_i] + "] " + partB.opinions[char_i][0]
            #table.rows[char_i].cells[1].left_indent = 21
            #table.rows[char_i].cells[1].first_line_indent = -21
            #p = table.rows[char_i].cells[1].add_paragraph("paragraph")
            #p.paragraph_format.left_indent = 0
            #p.paragraph_format.first_line_indent =-20
    
    elif partB.meta.test == 'NEEP-II-B2':
        dw.write_paragraph('Read the following text and answer the questions by choosing the most suitable subheading from the list A–G for each of the numbered paragraphs (41–45). There are two extra subheadings which you do not need to use. Mark your answers on the ANSWER SHEET. (10 points)', first_line_indent=0)
        dw.write_paragraph('')
        
        for i in range(0, 7):
            dw.write_paragraph('[' + labels[i] + '] ' + partB.opinions[i][0], left_indent=21, first_line_indent=-21)
        
        qb_index = 41
        for para in partB.article_st:
            text = ' '.join(para)
            #dw.write_paragraph(text)
            if re.match(r'\s*\#', text):
                dw.write_paragraph(str(qb_index) + '. ' + '_'*30)
                qb_index += 1
            else:
                dw.write_paragraph(text)   # write article paragraph.
    else:
        print("Error: Wrong Part B type.")
    #dw.write_paragraph(''.join(paper.partB.answers))
    #dw.write_paragraph('')
    dw.add_page_break()

def paper_write_tsl(dw, tsl):
    dw.write_paragraph('Section III  Translation', font_size=14, align='CENTER')
    dw.write_paragraph('46. Directions:', bold=True, space_after=5, first_line_indent=0)
    dw.write_paragraph('Translate the following text into Chinese. Write your translation on the ANSWER SHEET. (15 points)', first_line_indent=0)
    dw.write_paragraph('')
    for para in tsl.article_st:
        #dw.write_paragraph(trim(para))
        text = ' '.join(para)
        dw.write_paragraph(text)

def paper_write_writing(dw, writing, img_src):
    dw.write_paragraph('Section IV  Writing', font_size=14, align='CENTER')
    dw.write_paragraph('Part A', bold=True, first_line_indent=0, space_after=8)
    dw.write_paragraph('47. Directions:', bold=True, space_after=5, first_line_indent=0)

    for para in writing.l_direction:
        if re.match(r'\s*\*', para):
            break
        if re.match(r'\s*\d\)', para):
            dw.write_paragraph(para, left_indent=20)
        else:
            dw.write_paragraph(para)
    dw.write_paragraph('You should write about 100 words on the ANSWER SHEET.')
    p = dw.add_paragraph()
    p.add_run("Do not ")
    p.add_run("use your own name. Use “Li Ming” instead.")
    p = dw.add_paragraph()
    p.add_run("Do not ").bold = True
    p.add_run("write your address. (10 points)")

    dw.add_page_break()
    dw.write_paragraph('Part B', bold=True, first_line_indent=0, space_after=8)
    dw.write_paragraph('48. Directions:', bold=True, space_after=5, first_line_indent=0)

    for para in writing.e_direction:
        para = re.sub(r'\*\*', '', para)
        if re.match(r'\s*\d\)', para):
            dw.write_paragraph(para, left_indent=20)
        else:
            dw.write_paragraph(para)
    dw.add_picture(img_src, width=Inches(4.5))
    last_paragraph = dw.doc.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    
    
#---------------------------------------------------------------------------------------#
#                             Paper&Solution subs
#---------------------------------------------------------------------------------------#

def word_by_word(doc, text, words_dict):
    words = text.split()
    words_intpt_list = []
    #print(words)
    p = doc.add_paragraph('')
    
    for w in words:        
        word_and_punc = re.findall(r"[\w'\-\_]+|[\.\,!?;:\"\$]+|.", w)    # |. at the end makes sure nothing is left behind, including spaces.
        if word_and_punc:    # separate word and punctuation so be able to apply typesetting to words only.
            #print(word_and_punc)
            for w1 in word_and_punc:
                m1 = re.search(r"[\w'\-\_]+", w1)
                if m1:    # word
                    run = p.add_run(w1)
                    lemma = wn.morphy(w1)
                    if lemma in words_dict and lemma not in words_intpt_list:
                        run.bold = True
                        words_intpt_list.append(wn.morphy(w1))
                else:    # punctuation
                    p.add_run(w1)
                    
            #if w == '② ':
                #print(':' + w1 + ':')
                #print(word_and_punc)
                #pass
        else:    # none word and punctuation characters.
            p.add_run(w)
            
        
        p.add_run(' ')    # add space if 1) only a word; 2) word and punc, like:  why?  known.
        
    return words_intpt_list

def annotate_article(dw, article_st, translation_st, words_dict):
    #return
    for  para_i in range(0, len(article_st)):
    
        if re.match(r'^\s*\#\#', article_st[para_i][0]):
            continue    # partB titles like:  ## 41. The subtitle
            
        table = dw.add_table(rows=1, cols=2)
        text_en = ' '.join( [circled_number[stce_i]+' '+article_st[para_i][stce_i] for stce_i in range(0, len(article_st[para_i]))] )
        text_cn = ' '.join( [circled_number[stce_i]+' '+translation_st[para_i][stce_i] for stce_i in range(0, len(translation_st[para_i]))] )
        #print(text_cn)
        words_intpt_list = word_by_word(table.rows[0].cells[0], text_en, words_dict)
        table.rows[0].cells[1].add_paragraph(text_cn)
        
        if words_intpt_list:
            dw.write_paragraph('「词汇溯源」', align='LEFT', bold=True, left_indent=21, first_line_indent=-21)
            for w in words_intpt_list:
                p = dw.write_paragraph(w, bold=True, left_indent=0)
                if words_dict[w].phonetic != '':
                    p.add_run(' [' + words_dict[w].phonetic + ']')
            
                if words_dict[w].intpt != '':
                    p = dw.write_paragraph('[释义] ' + words_dict[w].intpt, font_size=10, align='LEFT', bold=False, left_indent=5, first_line_indent=21)
                    
                if words_dict[w].construct != '':
                    p = dw.write_paragraph('[构词] ' + '[' + words_dict[w].construct + ']', font_size=10, align='LEFT', bold=False, left_indent=5, first_line_indent=21)
                
                if words_dict[w].same_roots:
                    same_root_words = '    '.join(words_dict[w].same_roots)
                    p = dw.write_paragraph('[同根词] ' + same_root_words, font_size=10, align='LEFT', bold=False, left_indent=5, first_line_indent=21)
            
                if words_dict[w].syn:
                    synonyms = '    '.join(words_dict[w].syn)
                    p = dw.write_paragraph('[同义词] ' + synonyms, font_size=10, align='LEFT', bold=False, left_indent=5, first_line_indent=21)
            
                if words_dict[w].ant:
                    antonyms = '    '.join(words_dict[w].ant)
                    p = dw.write_paragraph('[反义词] ' + antonyms, font_size=10, align='LEFT', bold=False, left_indent=5, first_line_indent=21)
            
                del words_dict[w]

def paper_solution_write_closing(dw, closing, words_dict):
    dw.write_paragraph('Section I  Use of English', font_size=14, align='CENTER')
    dw.write_paragraph('')
    
    #dw.write_paragraph('Directions:', bold=True, space_after=5, first_line_indent=0)
    #dw.write_paragraph('Read the following text. Choose the best word(s) for each numbered blank and mark A, B, C or D on the ANSWER SHEET. (10 points)', first_line_indent=0)
    #dw.write_paragraph('')
    dw.write_paragraph('解析人：' + closing.meta.maker, bold=True)
    annotate_article(dw, closing.article_st, closing.translation_st, words_dict)
    
    #dw.add_page_break()
    
    for qc_idx in range(1, 21):
        p = dw.write_paragraph(str(qc_idx) + '. ')
        p.add_run('[A] ' + closing.Q[qc_idx-1].A[0] + '\t')
        p.add_run('[B] ' + closing.Q[qc_idx-1].B[0] + '\t')
        p.add_run('[C] ' + closing.Q[qc_idx-1].C[0] + '\t')
        p.add_run('[D] ' + closing.Q[qc_idx-1].D[0])
    
        dw.write_paragraph('「答案」'+closing.Q[qc_idx-1].answer, font_size=12, align='LEFT', bold=True)
        dw.write_paragraph('「试题详解」', align='LEFT', bold=True)
        dw.write_paragraph(closing.Q[qc_idx-1].exp, align='LEFT')
    #dw.add_page_break()
    
def paper_solution_write_partA(dw, partA, words_dict):
    dw.write_paragraph('Section II  Reading Comprehension', font_size=14, align='CENTER')
    dw.write_paragraph('')
    dw.write_paragraph('Part A', bold=True, first_line_indent=0, space_after=8)

    qa_index = 21
    for text_i in range(0, 4):
        p = dw.write_paragraph('Text '+ str(text_i+1), space_after=20, align='CENTER', bold=True)
        #p.paragraph_format.space_before = Pt(0)
        dw.write_paragraph('解析人：' + partA[text_i].meta.maker, bold=True)
        annotate_article(dw, partA[text_i].article_st, partA[text_i].translation_st, words_dict)
        #dw.add_page_break()  # PAGE BREAK AFTER PART A ARTICLE

        dw.write_paragraph('「习题解析」', font_size=12, align='LEFT', bold=True, first_line_indent=-21)
        for question_i in range(0, 5):
            Question = partA[text_i].Q[question_i]
            table = dw.add_table(rows = 5, cols = 2, style='Table Grid')
            table.rows[0].cells[0].text = str(qa_index) + '. ' + Question.question[0]
            table.rows[0].cells[1].text = str(qa_index) + '. ' + Question.question[1]
            rows_i = 1
            for data in (Question.A, Question.B, Question.C, Question.D):
                table.rows[rows_i].cells[0].text = '[' + labels[rows_i-1] + '] ' + data[0]
                table.rows[rows_i].cells[1].text = '[' + labels[rows_i-1] + '] ' + data[1]
                rows_i += 1
            dw.write_paragraph('')
        
            type = Question.type
            answer = Question.answer
       
            dw.write_paragraph('「试题类型」 ' + type_names[type], align='LEFT', bold=True, left_indent=21, first_line_indent=-21)
            dw.write_paragraph('「答案」 ' + answer, align='LEFT', bold=True, left_indent=21, first_line_indent=-21)
        
            p = dw.write_paragraph('「题目概览」', align='LEFT', bold=True, left_indent=10, first_line_indent=-10)
            p.add_run(Question.question[2])
            #print(Question.question[1])

            p = dw.write_paragraph('「选项解析」', align='LEFT', bold=True, left_indent=10, first_line_indent=-10)
            p.add_run('[A] ' + Question.A[2])
        
            dw.write_paragraph('[B] ' + Question.B[2], align='LEFT', left_indent=21, first_line_indent=-21)
            dw.write_paragraph('[C] ' + Question.C[2], align='LEFT', left_indent=21, first_line_indent=-21)
            dw.write_paragraph('[D] ' + Question.D[2], align='LEFT', left_indent=21, first_line_indent=-21)
            #print(Question.D[1])
            qa_index += 1

        #dw.add_page_break()  # PAGE BREAK AFTER PART A QUESTIONS

def paper_solution_write_partB(dw, partB, words_dict):

    dw.write_paragraph('Part B', bold=True, first_line_indent=0, space_after=8)
    dw.write_paragraph('解析人：' + partB.meta.maker, bold=True)
    if partB.meta.test == 'NEEP-II-B1':
        annotate_article(dw, partB.article_st, partB.translation_st, words_dict)

        dw.write_paragraph('')
        table = dw.add_table(rows=7, cols=2, style='Table Grid')
        qb_index = 41
        for char_i in range(1, 6):
            table.rows[char_i].cells[0].text = str(qb_index) + ". " +partB.characters[char_i-1]
            table.rows[char_i].cells[0].left_indent = 21
            table.rows[char_i].cells[0].first_line_indent = -21
            qb_index += 1

        for char_i in range(0, 7):
            table.rows[char_i].cells[1].text = "[" + labels[char_i] + "] " + partB.opinions[char_i][0]
            #table.rows[char_i].cells[1].left_indent = 21
            #table.rows[char_i].cells[1].first_line_indent = -21
            #p = table.rows[char_i].cells[1].add_paragraph("paragraph")
            #p.paragraph_format.left_indent = 0
            #p.paragraph_format.first_line_indent =-20
            
        dw.write_paragraph('「选项解析」', bold=True, align='LEFT', left_indent=10, first_line_indent=-10)
        qb_index = 41
        for title_i in range(0, 5):
            dw.write_paragraph(str(qb_index) + '. ' + partB.answers[title_i] + '. ' + partB.intpts[title_i][0])
            if partB.intpts[title_i][1] != '':
                p = dw.write_paragraph('「干扰项设置」', bold=True, align='LEFT')
                p.add_run(partB.intpts[title_i][1])
            qb_index += 1
    
    elif partB.meta.test == 'NEEP-II-B2':
        annotate_article(dw, partB.article_st, partB.translation_st, words_dict)
        
        dw.write_paragraph('「选项解析」', bold=True, align='LEFT', left_indent=10, first_line_indent=-10)
        qb_index = 41
        for title_i in range(0, 5):
            dw.write_paragraph(str(qb_index) + '. ' + partB.answers[title_i] + '. ' + partB.intpts[title_i][0])
            if partB.intpts[title_i][1] != '':
                p = dw.write_paragraph('「干扰项设置」', bold=True, align='LEFT')
                p.add_run(partB.intpts[title_i][1])
            qb_index += 1
            
    else:
        print("Error: Wrong Part B type.")
    #dw.write_paragraph(''.join(paper.partB.answers))
    #dw.write_paragraph('')
    #dw.add_page_break()
    
def paper_solution_write_tsl(dw, tsl, words_dict):
    dw.write_paragraph('Section III  Translation', font_size=14, align='CENTER')
    dw.write_paragraph('')
    
    dw.write_paragraph('解析人：' + tsl.meta.maker, bold=True)
    for para in tsl.article_st:
        #dw.write_paragraph(trim(para))
        text = ' '.join(para)
        #dw.write_paragraph(text)
    dw.write_paragraph('「参考译文」', bold=True, align='LEFT', left_indent=10, first_line_indent=-10)
    for para in tsl.translation_st:
        #dw.write_paragraph(trim(para))
        text = ' '.join(para)
        #dw.write_paragraph(text)
    annotate_article(dw, tsl.article_st, tsl.translation_st, words_dict)

    dw.write_paragraph('译者：' + tsl.meta.maker, align='RIGHT')

def paper_solution_write_writing(dw, writing):
    dw.write_paragraph('Section IV  Writing', font_size=14, align='CENTER')

    p = dw.write_paragraph('「写作思路」', bold=True, align='LEFT', left_indent=10, first_line_indent=-10)
    for para in writing.l_explanation:
        dw.write_paragraph(para)

    p = dw.write_paragraph('「参考范文」', bold=True, align='LEFT', left_indent=10, first_line_indent=-10)
    for para in writing.l_model:
        if zp.regex(re.match, r'\s*\<align\:\s*(.+)\s*\>\s*(.+)', para, 0, zp.r):    # <align:no-indent> Dear Professor Smith,
            align_str = zp.r.m.group(1)
            content = zp.r.m.group(2)
            if align_str == 'no-indent':
                dw.write_paragraph(content, left_indent=10, first_line_indent=-10)
            elif align_str == 'right':
                dw.write_paragraph(content, align='RIGHT')
            elif align_str == 'center':
                dw.write_paragraph(content, align='CENTER')
            else:  # useless by now
                dw.write_paragraph(content)
        else:
            dw.write_paragraph(para)

    p = dw.write_paragraph('「译文」', bold=True, align='LEFT', left_indent=10, first_line_indent=-10)
    for para in writing.l_translation:
        if zp.regex(re.match, r'\s*\<align\:\s*(.+)\s*\>\s*(.+)', para, 0, zp.r):    # <align:no-indent> Dear Professor Smith,
            align_str = zp.r.m.group(1)
            content = zp.r.m.group(2)
            if align_str == 'no-indent':
                dw.write_paragraph(content, left_indent=10, first_line_indent=-10)
            elif align_str == 'right':
                dw.write_paragraph(content, align='RIGHT')
            elif align_str == 'center':
                dw.write_paragraph(content, align='CENTER')
            else:  # useless by now
                dw.write_paragraph(content)
        else:
            dw.write_paragraph(para)


    dw.write_paragraph('')
    dw.write_paragraph('Part B', bold=True, first_line_indent=0, space_after=8)

    p = dw.write_paragraph('「写作思路」', bold=True, align='LEFT', left_indent=10, first_line_indent=-10)
    for para in writing.e_explanation:
        dw.write_paragraph(para)

    p = dw.write_paragraph('「参考范文」', bold=True, align='LEFT', left_indent=10, first_line_indent=-10)
    for para in writing.e_model:
        dw.write_paragraph(para)

    p = dw.write_paragraph('「译文」', bold=True, align='LEFT', left_indent=10, first_line_indent=-10)
    for para in writing.e_translation:
        dw.write_paragraph(para)


#---------------------------------------------------------------------------------------#
#                             Paper&Solution subs -- simple version
#---------------------------------------------------------------------------------------#

def paper_solution_simple_write_closing(dw, closing, words_dict):
    dw.write_paragraph('Section I  Use of English', font_size=14, align='CENTER')
    dw.write_paragraph('')
    
    #dw.write_paragraph('Directions:', bold=True, space_after=5, first_line_indent=0)
    #dw.write_paragraph('Read the following text. Choose the best word(s) for each numbered blank and mark A, B, C or D on the ANSWER SHEET. (10 points)', first_line_indent=0)
    #dw.write_paragraph('')
    # dw.write_paragraph('解析人：' + closing.meta.maker, bold=True)
    # annotate_article(dw, closing.article_st, closing.translation_st, words_dict)
    
    #dw.add_page_break()
    
    for qc_idx in range(1, 21):
        dw.write_paragraph(str(qc_idx) + '. ' + '「答案」'+closing.Q[qc_idx-1].answer, font_size=12, align='LEFT', bold=False)
        p = dw.write_paragraph('「试题详解」', align='LEFT', bold=True)
        p.add_run(closing.Q[qc_idx-1].exp)
    #dw.add_page_break()
    
def paper_solution_simple_write_partA(dw, partA, words_dict):
    dw.write_paragraph('Section II  Reading Comprehension', font_size=14, align='CENTER')
    dw.write_paragraph('')
    dw.write_paragraph('Part A', bold=True, first_line_indent=0, space_after=8)

    qa_index = 21
    for text_i in range(0, 4):
        p = dw.write_paragraph('Text '+ str(text_i+1), space_after=20, align='CENTER', bold=True)
        #p.paragraph_format.space_before = Pt(0)
        # dw.write_paragraph('解析人：' + partA[text_i].meta.maker, bold=True)
        # annotate_article(dw, partA[text_i].article_st, partA[text_i].translation_st, words_dict)
        # dw.add_page_break()  # PAGE BREAK AFTER PART A ARTICLE

        dw.write_paragraph('「习题解析」', font_size=12, align='LEFT', bold=True, first_line_indent=-21)
        for question_i in range(0, 5):
            try:
                Question = partA[text_i].Q[question_i]
                #''' # Uncomment this if you want a table of questions & its translations
                table = dw.add_table(rows = 5, cols = 2, style='Table Grid')
                table.rows[0].cells[0].text = str(qa_index) + '. ' + Question.question[0]
                table.rows[0].cells[1].text = str(qa_index) + '. ' + Question.question[1]
                rows_i = 1
                for data in (Question.A, Question.B, Question.C, Question.D):
                    try:
                        table.rows[rows_i].cells[0].text = '[' + labels[rows_i-1] + '] ' + data[0]
                        table.rows[rows_i].cells[1].text = '[' + labels[rows_i-1] + '] ' + data[1]
                    except Exception as e:
                        print(e)
                        print('paper_solution_simple_write_partA: ', data[0], data[1])
                    rows_i += 1
                dw.write_paragraph('')
                #'''
            
                type = Question.type
                answer = Question.answer
           
                dw.write_paragraph(str(qa_index) + '. ' + '「试题类型」 ' + type_names[type], align='LEFT', bold=True, left_indent=21, first_line_indent=-21)
                dw.write_paragraph('「答案」 ' + answer, align='LEFT', bold=True, left_indent=21, first_line_indent=-21)
            
                p = dw.write_paragraph('「题目概览」', align='LEFT', bold=True, left_indent=10, first_line_indent=-10)
                p.add_run(Question.question[2])
                #print(Question.question[1])

                p = dw.write_paragraph('「选项解析」', align='LEFT', bold=True, left_indent=10, first_line_indent=-10)
                p.add_run('[A] ' + Question.A[2])
            
                dw.write_paragraph('[B] ' + Question.B[2], align='LEFT', left_indent=21, first_line_indent=-21)
                dw.write_paragraph('[C] ' + Question.C[2], align='LEFT', left_indent=21, first_line_indent=-21)
                dw.write_paragraph('[D] ' + Question.D[2], align='LEFT', left_indent=21, first_line_indent=-21)
                #print(Question.D[1])
                qa_index += 1
            except Exception as e:
                print(e)
                print('paper_solution_simple_write_partA: error at Text ', text_i+1, ',question ', question_i, '\n')
        dw.write_paragraph('「参考译文」', bold=True, align='LEFT', left_indent=10, first_line_indent=-10)
        for para in partA[text_i].translation:
            dw.write_paragraph(para)

        #dw.add_page_break()  # PAGE BREAK AFTER PART A QUESTIONS

def paper_solution_simple_write_partB(dw, partB, words_dict):

    dw.write_paragraph('Part B', bold=True, first_line_indent=0, space_after=8)
    # dw.write_paragraph('解析人：' + partB.meta.maker, bold=True)
    if partB.meta.test == 'NEEP-II-B1':   
        dw.write_paragraph('「选项解析」', bold=True, align='LEFT', left_indent=10, first_line_indent=-10)
        qb_index = 41
        for title_i in range(0, 5):
            dw.write_paragraph(str(qb_index) + '. ' + partB.answers[title_i] + '. ' + partB.intpts[title_i][0])
            if partB.intpts[title_i][1] != '':
                p = dw.write_paragraph('「干扰项设置」', bold=True, align='LEFT')
                p.add_run(partB.intpts[title_i][1])
            qb_index += 1
    
    elif partB.meta.test == 'NEEP-II-B2':
        # annotate_article(dw, partB.article_st, partB.translation_st, words_dict)
        
        dw.write_paragraph('「选项解析」', bold=True, align='LEFT', left_indent=10, first_line_indent=-10)
        qb_index = 41
        for title_i in range(0, 5):
            dw.write_paragraph(str(qb_index) + '. ' + partB.answers[title_i] + '. ' + partB.intpts[title_i][0])
            if partB.intpts[title_i][1] != '':
                p = dw.write_paragraph('「干扰项设置」', bold=True, align='LEFT')
                p.add_run(partB.intpts[title_i][1])
            qb_index += 1
            
    else:
        print("Error: Wrong Part B type.")
    #dw.write_paragraph(''.join(paper.partB.answers))
    #dw.write_paragraph('')
    #dw.add_page_break()
    
def paper_solution_simple_write_tsl(dw, tsl, words_dict):
    dw.write_paragraph('Section III  Translation', font_size=14, align='CENTER')
    dw.write_paragraph('')
    
    # dw.write_paragraph('解析人：' + tsl.meta.maker, bold=True)
    dw.write_paragraph('「参考译文」', bold=True, align='LEFT', left_indent=10, first_line_indent=-10)
    for para in tsl.translation_st:
        text = ' '.join(para)
        dw.write_paragraph(text)

    # dw.write_paragraph('译者：' + tsl.meta.maker, align='RIGHT')

def paper_solution_simple_write_writing(dw, writing):
    dw.write_paragraph('Section IV  Writing', font_size=14, align='CENTER')

    p = dw.write_paragraph('「写作思路」', bold=True, align='LEFT', left_indent=10, first_line_indent=-10)
    for para in writing.l_explanation:
        dw.write_paragraph(para)

    p = dw.write_paragraph('「参考范文」', bold=True, align='LEFT', left_indent=10, first_line_indent=-10)
    for para in writing.l_model:
        if zp.regex(re.match, r'\s*\<align\:\s*(.+)\s*\>\s*(.+)', para, 0, zp.r):    # <align:no-indent> Dear Professor Smith,
            align_str = zp.r.m.group(1)
            content = zp.r.m.group(2)
            if align_str == 'no-indent':
                dw.write_paragraph(content, left_indent=10, first_line_indent=-10)
            elif align_str == 'right':
                dw.write_paragraph(content, align='RIGHT')
            elif align_str == 'center':
                dw.write_paragraph(content, align='CENTER')
            else:  # useless by now
                dw.write_paragraph(content)
        else:
            dw.write_paragraph(para)

    p = dw.write_paragraph('「译文」', bold=True, align='LEFT', left_indent=10, first_line_indent=-10)
    for para in writing.l_translation:
        if zp.regex(re.match, r'\s*\<align\:\s*(.+)\s*\>\s*(.+)', para, 0, zp.r):    # <align:no-indent> Dear Professor Smith,
            align_str = zp.r.m.group(1)
            content = zp.r.m.group(2)
            if align_str == 'no-indent':
                dw.write_paragraph(content, left_indent=10, first_line_indent=-10)
            elif align_str == 'right':
                dw.write_paragraph(content, align='RIGHT')
            elif align_str == 'center':
                dw.write_paragraph(content, align='CENTER')
            else:  # useless by now
                dw.write_paragraph(content)
        else:
            dw.write_paragraph(para)


    dw.write_paragraph('')
    dw.write_paragraph('Part B', bold=True, first_line_indent=0, space_after=8)

    p = dw.write_paragraph('「写作思路」', bold=True, align='LEFT', left_indent=10, first_line_indent=-10)
    for para in writing.e_explanation:
        dw.write_paragraph(para)

    p = dw.write_paragraph('「参考范文」', bold=True, align='LEFT', left_indent=10, first_line_indent=-10)
    for para in writing.e_model:
        dw.write_paragraph(para)

    p = dw.write_paragraph('「译文」', bold=True, align='LEFT', left_indent=10, first_line_indent=-10)
    for para in writing.e_translation:
        dw.write_paragraph(para)


