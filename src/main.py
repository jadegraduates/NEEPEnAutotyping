# -*- coding: utf-8 -*-
import os
import re
import random
import glob
import argparse
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE
import pandas as pd
from tqdm import tqdm

import word_structure as wds
from docx_wrapper import DocxWrapper
import main_subs as main_sub

pbar = tqdm(total=100)
#---------------------------------------------------------------------------------------#
#                             Loading Files
#---------------------------------------------------------------------------------------#
'''
main.py & main_subs.py is used to generate Exam paper & its solution file & the info file.
while 
main_gaofen.py & main_gaofen_subs.py is to generate Gaofen book.

'''

parser = argparse.ArgumentParser(description='NEEP English II paper Autotyping')
parser.add_argument('-d', '--directory', metavar='', required=True, help='required, directory for yaml surce files')
parser.add_argument('-n', '--name', metavar='', required=True, help="required, name prefix for output files'")
parser.add_argument('-t', '--title', metavar='', default='', required=False, help='paper title suffix, default Empty')


args = parser.parse_args()
# print(args)
# exit()

words_dict = wds.readWords(os.path.join('..', 'inputs', 'words_marked.txt'))
#stce_pool = sts.read_sentence('../input/gaofen/sentence_struct')

paper_name = args.name
source_dir = args.directory
paper_index = args.title

out_folder = source_dir

pbar.set_description('Loading source files from ' + source_dir + '...')
paper_files = {}

paper_files['Closing'] = glob.glob(os.path.join(source_dir, '*Closing*.yaml'))[0]
paper_files['PartA1'] = glob.glob(os.path.join(source_dir, '*Text1*.yaml'))[0]
paper_files['PartA2'] = glob.glob(os.path.join(source_dir, '*Text2*.yaml'))[0]
paper_files['PartA3'] = glob.glob(os.path.join(source_dir, '*Text3*.yaml'))[0]
paper_files['PartA4'] = glob.glob(os.path.join(source_dir, '*Text4*.yaml'))[0]
paper_files['PartB'] = glob.glob(os.path.join(source_dir, '*PartB*.yaml'))[0]
paper_files['Trans'] = glob.glob(os.path.join(source_dir, '*Trans*.yaml'))[0]
paper_files['Writing'] = glob.glob(os.path.join(source_dir, '*Writing*.yaml'))[0]
paper_files['img'] = ( glob.glob(os.path.join(source_dir, '*.png')) + glob.glob(os.path.join(source_dir, '*.jpg')) )[0]
img_src = paper_files['img']

paper = main_sub.paper_load_from_files(paper_files)

template_docx = os.path.join('..', 'inputs', 'template.docx')

#---------------------------------------------------------------------------------------#
#                             Generating  Paper
#---------------------------------------------------------------------------------------#
# pbar.update(10)
tqdm.write('Files loaded from ' + source_dir)
pbar.set_description('Generating paper ...')

dw = DocxWrapper(template_docx)

main_sub.paper_write_head_page(dw, paper_index)

main_sub.paper_write_closing(dw, paper.closing)

main_sub.paper_write_partA(dw, paper.partA)

#print(paper.partB.meta)
main_sub.paper_write_partB(dw, paper.partB)

main_sub.paper_write_tsl(dw, paper.tsl)

main_sub.paper_write_writing(dw, paper.writing, img_src)

dw.save( os.path.join(out_folder, paper_name + paper_index + '.docx') )


#---------------------------------------------------------------------------------------#
#                             Generating Paper PLUS Solution(Simple version)
#---------------------------------------------------------------------------------------#
pbar.update(20)
tqdm.write('Paper generated.')
pbar.set_description('Generating solution(simplified) ...')

dw = DocxWrapper(template_docx)

dw.write_paragraph('2019年全国硕士研究生入学统一考试', bold=True, align='CENTER', space_after=10, font_size=18)
dw.write_paragraph('英语（二）试卷' + paper_index + '解析', bold=True, align='CENTER', space_after=10, font_size=18)

main_sub.paper_solution_simple_write_closing(dw, paper.closing, words_dict)

main_sub.paper_solution_simple_write_partA(dw, paper.partA, words_dict)

main_sub.paper_solution_simple_write_partB(dw, paper.partB, words_dict)

main_sub.paper_solution_simple_write_tsl(dw, paper.tsl, words_dict)

main_sub.paper_solution_simple_write_writing(dw, paper.writing)


dw.save( os.path.join(out_folder, paper_name + paper_index + '-解析.docx') )


#---------------------------------------------------------------------------------------#
#                             Generating Paper PLUS Solution (Sophisticated version)
#---------------------------------------------------------------------------------------#
pbar.update(30)
tqdm.write('Solution(simplified) generated.')

if False:
  pbar.set_description('Generating solution(sophisticated) ...')
  dw = DocxWrapper(template_docx)

  dw.write_paragraph('2019年全国硕士研究生入学统一考试', bold=True, align='CENTER', space_after=10, font_size=18)
  dw.write_paragraph('英语（二）试卷' + paper_index + '解析', bold=True, align='CENTER', space_after=10, font_size=18)

  main_sub.paper_solution_write_closing(dw, paper.closing, words_dict)

  main_sub.paper_solution_write_partA(dw, paper.partA, words_dict)

  main_sub.paper_solution_write_partB(dw, paper.partB, words_dict)

  main_sub.paper_solution_write_tsl(dw, paper.tsl, words_dict)

  main_sub.paper_solution_write_writing(dw, paper.writing)


  dw.save( os.path.join(out_folder, paper_name + paper_index + '-解析（详尽版）.docx') )

  tqdm.write('Solution(sophisticated) generated.')
#---------------------------------------------------------------------------------------#
#                             Generating Bare Answers
#---------------------------------------------------------------------------------------#
pbar.update(30)
pbar.set_description('Generating bare answer ...')
# doc_resolution = Document(template_docx)
# doc_resolution.add_heading('一、参考答案', level=1)
# doc_resolution.add_heading('Section I Use of English', level=2)

dw = DocxWrapper(template_docx)
dw.add_heading('参考答案', level=3)
dw.write_paragraph('Section I Use of English', font_size=14, align='CENTER')

p = dw.write_paragraph('1-5. ')
for i in range(0, 5):
    p.add_run(paper.closing.Q[i].answer)

p.add_run('     6-10. ', '')
for i in range(5, 10):
    p.add_run(paper.closing.Q[i].answer)
    
p.add_run('     11-15. ', '')
for i in range(10, 15):
    p.add_run(paper.closing.Q[i].answer)

p.add_run('     16-20. ', '')
for i in range(15, 20):
    p.add_run(paper.closing.Q[i].answer)

dw.write_paragraph('Section II Part A', font_size=14, align='CENTER')
answers = ''
qa_index = 21
for text_i in range(0, 4):
    qa_top = qa_index + 4
    num_label = str(qa_index) + '-' + str(qa_top) + '. '
    qa_index += 5
    answers += num_label
    for question_i in range(0, 5):
        answers += paper.partA[text_i].Q[question_i].answer
    answers += '  '
dw.write_paragraph(answers)

dw.write_paragraph('Section II Part B', font_size=14, align='CENTER')
qb_index = 41
answers = ''.join(paper.partB.answers)
answers = '41-45. ' + answers
dw.write_paragraph(answers)

main_sub.paper_solution_simple_write_tsl(dw, paper.tsl, words_dict)

dw.save( os.path.join(out_folder, paper_name + paper_index + '-答案.docx') )



#---------------------------------------------------------------------------------------#
#                             Generating Paper Info
#---------------------------------------------------------------------------------------#
pbar.update(20)
tqdm.write('Bare answer generated.')
pbar.set_description('Generating paper info xlsx ...')

'''---------------- Paper Info statistics ----------------'''
keys_to_save = ['journal', 'date', 'type', 'words', 'title', 'author', 'maker', 'time', 'url']

df_info = pd.DataFrame(columns=keys_to_save)
for meta in (paper.closing.meta, paper.partA[0].meta, paper.partA[1].meta, paper.partA[2].meta, paper.partA[3].meta, paper.partB.meta, paper.tsl.meta):
  df_info = df_info.append({k:meta[k] for k in keys_to_save}, ignore_index=True)
df_info.rename(columns={'journal':'期刊', 'date':'发表日期', 'type':'文章类型', 'words':'节选字数', 'title':'标题', 'author':'作者', 'maker':'命题人', 'time':'命题日期', 'url':'链接'}, inplace=True)

# df_info = pd.concat([pd.Series(['Closing', 'Text 1', 'Text 2', 'Text 3', 'Text 4', 'Part B', 'Trans']), df_info])
df_info.insert(0, '题型', ['Closing', 'Text 1', 'Text 2', 'Text 3', 'Text 4', 'Part B', 'Trans'], True)
# df_info.to_excel(os.path.join(out_folder, paper_name + paper_index + '-信息.xlsx'), sheet_name='选文信息', index=False)

'''---------------- Part A type statistics ----------------'''
df_type_num = pd.DataFrame(columns=['info', 'com', 'vip', 'wd', 'eg', 'infer', 'att', 'pp'])
for partA in (paper.partA[0], paper.partA[1], paper.partA[2], paper.partA[3]):
  dict_types = {}
  for q in partA.Q:
    dict_types[q.type] = dict_types[q.type] + 1 if q.type in dict_types else 1
  
  df_type_num = df_type_num.append(dict_types, ignore_index=True)

df_type_num = df_type_num.append(df_type_num.agg(['sum']))
df_type_num['合计'] = df_type_num.agg('sum', axis="columns")

'''---------------- Closing type statistics ----------------'''
df_type_num.insert(0, '文章', ['Text 1', 'Text 2', 'Text 3', 'Text 4', '合计'])

df_closing_type_num = pd.DataFrame()
dict_types = {}
for q in paper.closing.Q:
  dict_types[q.type] = dict_types[q.type] + 1 if q.type in dict_types else 1
df_closing_type_num = df_closing_type_num.append(dict_types, ignore_index=True)
df_closing_type_num['合计'] = df_closing_type_num.iloc[0, :].sum()
df_closing_type_num.insert(0, '文章', ['Closing'])

'''-------------------- Write to Excel --------------------'''
pd_excel_writer = pd.ExcelWriter(os.path.join(out_folder, paper_name + paper_index + '-信息.xlsx'), engine='xlsxwriter')
df_info.to_excel(pd_excel_writer, sheet_name='选文信息', index=False)
df_type_num.to_excel(pd_excel_writer, sheet_name='PartA题型统计', index=False)
df_closing_type_num.to_excel(pd_excel_writer, sheet_name='完形题型统计', index=False)
pd_excel_writer.save()

tqdm.write('Paper info xlsx generated.')
tqdm.write('Autotyping completed, check out files at ' + out_folder)