
#---------------------------------------------------------------------------------------#
#                             Generating DOCX
#---------------------------------------------------------------------------------------#

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
#from docx.enum.style import WD_STYLE

import re

from nltk.corpus import wordnet as wn



class DocxWrapper:
    def __init__(self, model_file):
        self.doc = Document(model_file)
        self.doc.styles['Normal'].font.name = 'Times New Roman'
        self.doc.styles['Normal'].font.size = Pt(10.5)
        self.doc.styles['Normal'].paragraph_format.space_before = Pt(0)
        self.doc.styles['Normal'].paragraph_format.space_after = Pt(0)
        # document.styles['Normal'].paragraph_format.left_indent = Pt(36)
        self.doc.styles['Normal'].paragraph_format.first_line_indent = Pt(21)
        self.doc.styles['Normal'].paragraph_format.line_spacing = 1.25
        
        
    def add_paragraph(self, *args, **kwargs):
        return self.doc.add_paragraph(*args, **kwargs)
        
    def write_u_list(self, text):
        return self.add_paragraph(text, style='List Bullet')
        
    def add_picture(self, *args, **kwargs):
        return self.doc.add_picture(*args, **kwargs)
        
    def add_heading(self, text, level):
        return self.doc.add_heading(text, level=level)
    
    def add_table(self, *args, **kwargs):
        return self.doc.add_table(*args, **kwargs)
    
    def write_paragraph(self, text, font_size=10.5, space_after=0, align='JUSTIFY', bold=False, left_indent=0, first_line_indent=21):
        p = self.add_paragraph()
        p.paragraph_format.left_indent = Pt(left_indent)
        p.paragraph_format.first_line_indent = Pt(first_line_indent)
        #print("left line indent: ", p.paragraph_format.left_indent)
        run = p.add_run(text)
        #run.font.name = 'Times New Roman'
        run.font.size = Pt(font_size)
        if align == 'JUSTIFY':
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif align == 'LEFT':
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif align == 'CENTER':
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'RIGHT':
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
        if bold:
            run.bold = True
        #if space_after:
        p.paragraph_format.space_after =  Pt(space_after)
        return p
    
    def add_page_break(self, *args, **kwargs):
        return self.doc.add_page_break(*args, **kwargs)
    
    def save(self, file_name):
        return self.doc.save(file_name)






def write_paragraph(document, prefix, text, font_size=10.5, space_after=0, align='JUSTIFY', bold=False, left_indent=0, first_line_indent=21):
    # text = trim(text)
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Pt(left_indent)
    p.paragraph_format.first_line_indent = Pt(first_line_indent)
    #print("left line indent: ", p.paragraph_format.left_indent)
    text = prefix + text
    run = p.add_run(text)
    #run.font.name = 'Times New Roman'
    run.font.size = Pt(font_size)
    if align == 'JUSTIFY':
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == 'LEFT':
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == 'CENTER':
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'RIGHT':
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    if bold:
        run.bold = True
    #if space_after:
    p.paragraph_format.space_after =  Pt(space_after)
    return p

def write_run(p, text, font_size=10.5, bold=False):
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    
    if bold:
        run.bold = True
    return p


def paragraph(doc, text, font_size=10.5, space_after=0, align='JUSTIFY', bold=False, left_indent=0, first_line_indent=21):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(left_indent)
    p.paragraph_format.first_line_indent = Pt(first_line_indent)
    
    #text = trim(text)
    
    run = p.add_run(text)
    #run.font.name = 'Times New Roman'
    run.font.size = Pt(font_size)
    if align == 'JUSTIFY':
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == 'LEFT':
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == 'CENTER':
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'RIGHT':
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    if bold:
        run.bold = True
    #if space_after:
    p.paragraph_format.space_after =  Pt(space_after)
    return p

def text_en(doc, text, l=0, f=21):
    p = paragraph(doc, text, font_size=10.5, space_after=0, align='JUSTIFY', bold=False, left_indent=l, first_line_indent=f)
    return p

def text_cn(doc, text, l=0, f=21):
    p = paragraph(doc, text, font_size=10.5, space_after=0, align='LEFT', bold=False, left_indent=l, first_line_indent=f)
    return p
    
def title_center(doc, text):
    p = paragraph(doc, text, font_size=12, align='CENTER', bold=True)
    return p

def title_1(doc, text):
    p = paragraph(doc, text, font_size=12, align='LEFT', bold=True)
    return p

def title_2(doc, text, l=21, f=-21):
    p = paragraph(doc, text, font_size=10, align='LEFT', bold=True, left_indent=l, first_line_indent=f)
    return p

def title_3(doc, text, l=10, f=-10):
    p = paragraph(doc, text, font_size=10, align='LEFT', bold=True, left_indent=l, first_line_indent=f)
    return p

def title_word(doc, text):
    p = paragraph(doc, text, font_size=10, align='LEFT', bold=True, left_indent=0, first_line_indent=21)
    return p

def title_word_2(doc, text):
    p = paragraph(doc, text, font_size=10, align='LEFT', bold=False, left_indent=5, first_line_indent=21)
    return p

def text_en_type(doc, text, words_dict, l=0, f=21):
    words = text.split()
    words_intpt_list = []
    #print(words)
    p = paragraph(doc, '', font_size=10.5, space_after=0, align='JUSTIFY', bold=False, left_indent=l, first_line_indent=f)
    
    p.add_run(circled_number[0] + ' ')
    circled_i = 1;
    for w in words:    # sentence separation mark like /2/
        m = re.match(r'\/\d*\/', w)
        if m:
            w = circled_number[circled_i]
            circled_i += 1
        
        word_and_punc = re.findall(r"[\w'\-\_]+|[\.\,!?;:\"\$]+|.", w)    # |. at the end makes sure nothing is left behind, including spaces.
        if word_and_punc:    # separate word and punctuation so be able to apply typesetting to words only.
            #print(word_and_punc)
            for w1 in word_and_punc:
                m1 = re.search(r"[\w'\-\_]+", w1)
                if m1:    # word
                    run = p.add_run(w1)
                    lemma = wn.morphy(w1)
                    if lemma in words_dict and lemma not in words_intpt_list:
                        run.bold = True
                        words_intpt_list.append(wn.morphy(w1))
                else:    # punctuation
                    p.add_run(w1)
                    
            #if w == '② ':
                #print(':' + w1 + ':')
                #print(word_and_punc)
                #pass
        else:    # none word and punctuation characters.
            p.add_run(w)
            
        
        p.add_run(' ')    # add space if 1) only a word; 2) word and punc, like:  why?  known.
        
    return words_intpt_list
        

def text_cn_type(doc, text, l=0, f=21):
    parts = re.split(r"\s*\/\d*\/\s*", text)
    p = paragraph(doc, '', font_size=10.5, space_after=0, align='LEFT', bold=False, left_indent=l, first_line_indent=f)
    
    circled_i = 0
    for pt in parts:
        
        p.add_run(circled_number[circled_i] + ' ')
        p.add_run(pt)
        circled_i += 1
   



'''
def text_en_type(doc, text, l=0, f=21):
    # words = text.split()
    words = re.findall(r"[\w'\-\_]+|[.,!?;]", text, re.UNICODE)
    print(words)
    p = paragraph(doc, '', font_size=10.5, space_after=0, align='JUSTIFY', bold=False, left_indent=l, first_line_indent=f)
    
    p.add_run(circled_number[0] + ' ')
    circled_i = 1;
    for w in words:
        m = re.match(r'\/\d+?\/', w)
        if m:
            w = circled_number[circled_i]
            circled_i += 1
        p.add_run(w)
        m = re.match(r"[.,!?;]", w)
        if not m:
            p.add_run(' ')

'''


